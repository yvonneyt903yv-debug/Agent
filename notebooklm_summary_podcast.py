#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import asyncio
import os
import sys
import json
import argparse
from datetime import datetime
from docx import Document
from pathlib import Path

# ==========================================
# 🌍 网络代理设置 (解决 ConnectError)
# ==========================================
# 请替换为你自己的代理地址，例如 Clash 默认是 7890
PROXY_URL = "http://127.0.0.1:7890" 

os.environ["HTTP_PROXY"] = PROXY_URL
os.environ["HTTPS_PROXY"] = PROXY_URL
# 某些库可能还需要这个
# 添加项目路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 导入 NotebookLM 工具
from notebook_tool import NotebookSkill

# ==========================================
# 🛑【关键修改】强制指定输出目录 (绝对路径)
# ==========================================
TARGET_OUTPUT_DIR = "/Users/yvonne/Documents/MyAgent/output/podcasts"

# 确保这个目录存在，不存在就创建
if not os.path.exists(TARGET_OUTPUT_DIR):
    os.makedirs(TARGET_OUTPUT_DIR)
    print(f"📁 已创建目录: {TARGET_OUTPUT_DIR}")
else:
    print(f"📂以此目录作为输出目标: {TARGET_OUTPUT_DIR}")

def read_document_content(file_path):
    """
    读取文档内容，支持 .docx 和 .md 文件
    :param file_path: 文件路径
    :return: 文档内容字符串
    """
    file_ext = Path(file_path).suffix.lower()

    try:
        if file_ext == '.docx':
            doc = Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            print(f"📖 已读取 Word 文档: {file_path}")
            return content
        elif file_ext == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"📖 已读取 Markdown 文档: {file_path}")
            return content
        else:
            print(f"⚠️ 不支持的文件格式: {file_ext}")
            return None
    except Exception as e:
        print(f"❌ 读取文���失败: {e}")
        return None

def save_summary_to_word(summary_text, source_file, full_file_path):
    """
    将总结内容保存为 Word 文档
    :param summary_text: 总结内容
    :param source_file: 源文件路径
    :param full_file_path: 完整的输出文件绝对路径
    """
    try:
        doc = Document()

        # 添加标题
        doc.add_heading('NotebookLM 文章总结', 0)

        # 添加元数据
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"来源文件: {source_file}")
        doc.add_paragraph("-" * 20)

        # 添加正文
        doc.add_heading('核心要点', level=1)
        doc.add_paragraph(summary_text)

        # 保存到指定绝对路径
        doc.save(full_file_path)
        print(f"📄 Word 文档已保存: {full_file_path}")
        return True
    except Exception as e:
        print(f"❌ Word 保存失败: {e}")
        return False

async def process_document_with_notebooklm(file_path, file_content):
    """
    使用 NotebookLM 处理单个文档
    :param file_path: 文件路径（用于命名）
    :param file_content: 文件内容
    """
    try:
        async with NotebookSkill() as skill:

            # 1. 创建笔记本
            file_name = Path(file_path).stem
            title = f"{file_name}_{datetime.now().strftime('%Y%m%d_%H%M')}"
            nb = await skill.create_notebook(title)
            nb_id = str(nb.id) if hasattr(nb, 'id') else str(nb)
            print(f"📝 笔记本 ID: {nb_id}")

            # 2. 上传内容
            print(f"📤 正在上传文档内容...")
            await skill.upload_text(nb_id, file_content, title=file_name)

            # 3. 获取总结并保存 Word
            print(f"🤖 正在请求 NotebookLM 总结...")
            summary = await skill.ask_question(nb_id, "请用中文总结这篇文章的核心要点，结构清晰。")

            word_full_path = None
            if summary:
                word_filename = f"summary_{file_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                word_full_path = os.path.join(TARGET_OUTPUT_DIR, word_filename)
                save_summary_to_word(summary, file_path, word_full_path)

            # 4. 生成播客
            print(f"🎙️ 正在生成播客...")
            podcast_full_path = None

            try:
                podcast_instructions = "Focus on key insights for a Chinese audience.确保输出为中文"
                generation_status = await skill.client.artifacts.generate_audio(
                    nb_id,
                    instructions=podcast_instructions
                )

                # 设置超时为 900 秒（15 分钟）
                print(f"⏳ 等待播客生成完成（最多等待 15 分钟）...")
                await skill.client.artifacts.wait_for_completion(nb_id, generation_status.task_id, timeout=900)

                # Keep the downloaded audio extension aligned with NotebookLM container format.
                podcast_filename = f"podcast_{file_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.m4a"
                podcast_full_path = os.path.join(TARGET_OUTPUT_DIR, podcast_filename)

                await skill.client.artifacts.download_audio(nb_id, podcast_full_path)
                print(f"🎧 播客已下载: {podcast_full_path}")

            except Exception as e:
                print(f"⚠️ 播客生成/下载失败: {e}")

            return {
                'success': True,
                'word_path': word_full_path,
                'podcast_path': podcast_full_path
            }

    except Exception as e:
        print(f"❌ 处理异常: {e}")
        return {'success': False, 'error': str(e)}
async def main():
    # 解析命令行参数
    parser = argparse.ArgumentParser(
        description='使用 NotebookLM 生成文档总结和播客',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  # 处理单个文件
  python3 notebooklm_summary_podcast.py article.docx
  python3 notebooklm_summary_podcast.py article.md

  # 批量处理目录中的所有文件
  python3 notebooklm_summary_podcast.py
        '''
    )
    parser.add_argument('file', nargs='?', help='要处理的 Word (.docx) 或 Markdown (.md) 文件路径')
    args = parser.parse_args()

    print("=" * 60)
    print(f"输出目录: {TARGET_OUTPUT_DIR}")
    print("=" * 60)

    # 单文件模式
    if args.file:
        file_path = args.file

        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return

        # 检查文件格式
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in ['.docx', '.md']:
            print(f"❌ 不支持的文件格式: {file_ext}")
            print("支持的格式: .docx, .md")
            return

        print(f"\n📄 单文件模式: {file_path}")

        # 读取文件内容
        content = read_document_content(file_path)
        if not content:
            print("❌ 无法读取文件内容")
            return

        # 处理文档
        result = await process_document_with_notebooklm(file_path, content)

        if result['success']:
            print("\n✅ 处理完成!")
            if result.get('word_path'):
                print(f"📄 总结文档: {result['word_path']}")
            if result.get('podcast_path'):
                print(f"🎧 播客文件: {result['podcast_path']}")
        else:
            print(f"\n❌ 处理失败: {result.get('error', '未知错误')}")

    # 批量处理模式
    else:
        source_dir = "/Users/yvonne/Documents/MyAgent/output/final_published"

        print(f"\n📁 批量处理模式")
        print(f"源目录: {source_dir}")

        if not os.path.exists(source_dir):
            print(f"❌ 源目录不存在: {source_dir}")
            return

        # 查找所有支持的文件
        files = [f for f in os.listdir(source_dir) if f.endswith(('.md', '.docx'))]

        if not files:
            print("⚠️ 没有找到 .md 或 .docx 文件")
            return

        print(f"找到 {len(files)} 个文件\n")

        for file_name in files:
            file_path = os.path.join(source_dir, file_name)
            print(f"\n{'='*60}")
            print(f"处理文件: {file_name}")
            print(f"{'='*60}")

            # 读取文件内容
            content = read_document_content(file_path)
            if not content:
                print(f"⚠️ 跳过文件: {file_name}")
                continue

            # 处理文档
            result = await process_document_with_notebooklm(file_path, content)

            if result['success']:
                print(f"✅ {file_name} 处理完成")
            else:
                print(f"❌ {file_name} 处理失败: {result.get('error', '未知错误')}")

        print(f"\n{'='*60}")
        print("🎉 批量处理完成!")
        print(f"{'='*60}")

if __name__ == "__main__":
    asyncio.run(main())
